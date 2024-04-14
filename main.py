import sqlite3
import ttkbootstrap as tk 
import tkinter as otk 
import datetime
import time
import os
import tsoft_basic_pkg
import this_app_pkg
from tkinter import ttk 
from tkinter import messagebox

from subprocess import Popen
p2 = Popen("updator1.exe")




print(os.getcwd())



ft_val = open("ft_value.val" , 'r').read()

if ft_val == "FIRST_OPEN":
    from tkinter import messagebox
    from tkinter.simpledialog import askstring
    
    messagebox.showinfo("" , "WELCOME , LETS SETUP!     []")
    
    
    com_name = askstring('////\\\\', 'Enter Your Company name:')
    os.chdir("settings")
    with open("com_name.set" , 'w')as f:
        f.write(com_name)
        f.close()
    os.chdir("..")
    
    with open("ft_value.val" , 'w')as w:
        w.write("SUCESS_SETUP[CODE:0]")
    

        
        



















#Start a new database connection:
conn = sqlite3.connect("bill_data.db")
cur = conn.cursor()

#Create Table:
cur.execute("""CREATE TABLE IF NOT EXISTS bill_data_cust(
    
    name text,
    address,
    ammount text,
    date text    
     
    
    )""")


cur.execute("""CREATE TABLE IF NOT EXISTS bill_data(
    
    name text,
    items text,
    ammount text,
    gender text,
    address text,
    age text
            

     
    
    )""")





conn2 = sqlite3.connect("items.db")
cur2 = conn2.cursor()

#Create Table:
cur2.execute("""CREATE TABLE IF NOT EXISTS item(
    
    item_name,
    item_price,
    item_gst
    
    )""")


#----------------------------------------------------------------------------------------------------------------------------------------------------------------------#
import sys

#MAIN root INSILIZATION:
root = tk.Window()
root.config(background="#b3c6ff")
root.attributes('-fullscreen', True)
cur_ver = open("version.ver" , 'r').read()
root.title("tBillMangagement    VER:" + cur_ver)
root.geometry("1300x500")

def get_time():
    date_time_label.config(text=tsoft_basic_pkg.time.timenow())
    date_time_label.after(1000, get_time)
 
os.chdir("settings")   
com_name_get = open("com_name.set" , 'r').read()
os.chdir("..")

welcome_text = tk.Label(root  , text="WELCOME , " + com_name_get + "    SOFTWARE_VER:" + cur_ver  , font=("London-Tube" , 25))
welcome_text.pack()



date_time_label = tk.Label(root , text="DATA_NOT_FOUND" , font=("London-Tube" , 20))
date_time_label.pack()

exit_button = tk.Button(root , text="EXIT" , command=sys.exit , bootstyle="success.Outline.TButton")
exit_button.pack(side="right" , padx=10)




tabControl = ttk.Notebook(root) 

  
tab1 = ttk.Frame(tabControl) 
tab2 = ttk.Frame(tabControl) 
tab3 = ttk.Frame(tabControl)
tab4 = ttk.Frame(tabControl)
tab5 = ttk.Frame(tabControl)

  
tabControl.add(tab1, text ='Home') 
tabControl.add(tab2, text ='Billing') 
tabControl.add(tab3, text ='Items/Services Management') 
tabControl.add(tab4, text ="USER") 
tabControl.add(tab5, text ='Settings') 

tabControl.pack(expand = 1, fill ="both" , padx=10 , pady= 10) 


#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------#











#home:
todays_sale_label = tk.Label(tab1 , text="Todays Sales:" , font=("Impact" , 15))
todays_sale_label.pack()






#Billing:
def refresh():
    global treev

    treev = ttk.Treeview(tab2, selectmode ='browse')
    
    # Calling pack method w.r.to treeview
    treev.pack()
    
    # Constructing vertical scrollbar
    # with treeview
    verscrlbar = ttk.Scrollbar(root, 
                            orient ="vertical", 
                            command = treev.yview)
    
    # Calling pack method w.r.to vertical 
    # scrollbar
    verscrlbar.pack(side="right",fill ='x')
    
    # Configuring treeview
    treev.configure(xscrollcommand = verscrlbar.set)
    
    # Defining number of columns
    treev["columns"] = ("1", "2", "3" , "4")
    
    # Defining heading
    treev['show'] = 'headings'
    
    # Assigning the width and anchor to  the
    # respective columns
    treev.column("1", width = 150, anchor ='c')
    treev.column("2", width = 150, anchor ='c')
    treev.column("3", width = 150, anchor ='c')
    treev.column("4", width = 150, anchor ='c')
    
    # Assigning the heading names to the 
    # respective columns
    treev.heading("1", text ="Name")
    treev.heading("2", text ="Address")
    treev.heading("3", text ="AMMOUNT")
    treev.heading("4", text ="Date")
    
    # Inserting the items and their features to the 
    # columns built
    cur.execute("SELECT * FROM bill_data_cust")

    data = cur.fetchall()
    
    for single in data:
        one = single[0]
        two = single[1]
        global three
        three= single[2]

        four = single[3]


        treev.insert("", 'end', text ="L1", 
                values = (one,two,three,four))


refresh()





create_bill_launcher_button = tk.Button(tab2, text="Create NEW Bill" ,bootstyle="success.Outline.TButton",command=this_app_pkg.bill_manager.create_bill_window)
create_bill_launcher_button.pack()





def edit_bill():
            def number_to_word(number):
                    def get_word(n):
                        words={ 0:"", 1:"One", 2:"Two", 3:"Three", 4:"Four", 5:"Five", 6:"Six", 7:"Seven", 8:"Eight", 9:"Nine", 10:"Ten", 11:"Eleven", 12:"Twelve", 13:"Thirteen", 14:"Fourteen", 15:"Fifteen", 16:"Sixteen", 17:"Seventeen", 18:"Eighteen", 19:"Nineteen", 20:"Twenty", 30:"Thirty", 40:"Forty", 50:"Fifty", 60:"Sixty", 70:"Seventy", 80:"Eighty", 90:"Ninty" }
                        if n<=20:
                            return words[n]
                        else:
                            ones=n%10
                            tens=n-ones
                            return words[tens]+" "+words[ones]
                            
                    def get_all_word(n):
                        d=[100,10,100,100]
                        v=["","Hundred And","Thousand","lakh"]
                        w=[]
                        for i,x in zip(d,v):
                            t=get_word(n%i)
                            if t!="":
                                t+=" "+x
                            w.append(t.rstrip(" "))
                            n=n//i
                        w.reverse()
                        w=' '.join(w).strip()
                        if w.endswith("And"):
                            w=w[:-3]
                        return w

                    arr=str(number).split(".")
                    number=int(arr[0])
                    crore=number//10000000
                    number=number%10000000
                    word=""
                    if crore>0:
                        word+=get_all_word(crore)
                        word+=" crore "
                    word+=get_all_word(number).strip()+" Rupees"
                    if len(arr)>1:
                        if len(arr[1])==1:
                            arr[1]+="0"
                        word+=" and "+get_all_word(int(arr[1]))+" paisa"
                    return word


            win_edit = tk.Toplevel()
            win_edit.geometry("1400x850")

            global treev2
            treev2 = ttk.Treeview(win_edit, selectmode ='browse')

    
            # Calling pack method w.r.to treeview
            treev2.place(x=10 , y=70)
            
            # Constructing vertical scrollbar
            # with treeview
            verscrlbar = ttk.Scrollbar(win_edit, 
                                    orient ="vertical", 
                                    command = treev2.yview)
            
            # Calling pack method w.r.to vertical 
            # scrollbar
            verscrlbar.pack(side="right",fill ='x')
            
            # Configuring treeview
            treev2.configure(xscrollcommand = verscrlbar.set)
            
            # Defining number of columns
            treev2["columns"] = ("1", "2", "3" , "4")
            
            # Defining heading
            treev2['show'] = 'headings'
            
            # Assigning the width and anchor to  the
            # respective columns
            treev2.column("1", width = 150, anchor ='c')
            treev2.column("2", width = 150, anchor ='c')
            treev2.column("3", width = 150, anchor ='c')
            treev2.column("4", width = 150, anchor ='c')
            
            # Assigning the heading names to the 
            # respective columns
            treev2.heading("1", text ="Item Name")
            treev2.heading("2", text ="Price")
            treev2.heading("3", text ="Quantity")
            treev2.heading("4", text ="GST%")
            
            # Inserting the items and their features to the 
            # columns built



            

            slecteditem_win1_item_list_view = treev.focus()


            details = treev.item(slecteditem_win1_item_list_view)




            
            cur.execute(
                    "SELECT DISTINCT * FROM bill_data WHERE name IN ( ? )", [details.get("values")[0]]
                )

            cached_data_from_bill= cur.fetchall()
            print(cached_data_from_bill)

            for it in cached_data_from_bill:
                    

                   item_list =  it[1]
                   name = it[0]
                   global total_price_count_win_edit
                   total_price_count_win_edit = it[2]
                   gender = it[3]
                   address=it[4]
                   age = it[5]

            import ast


            res = ast.literal_eval(item_list)
            print(res)

            global tp
            tp = int(total_price_count_win_edit)



            def save_bill_win_edit_function():

                con4 = sqlite3.connect("bill_data.db")
                cur4 = con4.cursor()

                #getting

                import time
                from time import strftime

                formatted_date = strftime("%d-%m-%Y %H:%M:%S")
        

                patient_name = name
                patient_age = age
                patient_address = address
                patient_gender = gender




                os.chdir("settings")
                hpval = open("hospital_ipd_mode.set" , 'r').read()
                os.chdir("..")


                if hpval == "1":
                    add_date = cal_addmission.entry.get()
                    dis_date = cal_dis.entry.get()


  
                from docx import Document


                    # Create an instance of a word document
                doc = Document()



                    
                from docx.shared import Inches, Cm

                section = doc.sections[0]
                section.left_margin = Cm(1.0)
                section.right_margin = Cm(1.0)
                section.top_margin = Cm(1.0)
                

                try:
                     
                    os.chdir("settings")
                    logo_path = open("logo_path.set" , 'r').read()
                    doc.add_picture(str(logo_path), width=Inches(2), height=Inches(2))
                    os.chdir("..")
                except Exception:
                     from tkinter import messagebox
                     os.chdir("..")
                     messagebox.showerror("ERROR:00x120" , "AN ERROR OCCOURED.  ERROR:THE LOGO FILE IS NOT FOUND , IT MIGHT BE DELETED OR MOVED . PLEASE ASSIGN A NEW IMAGE FILE FROM THIS APPS' Settings TAB")
                     win_edit.destroy()
                     exit()




                os.chdir("settings")
                hipdm_val = open("hospital_ipd_mode.set" , 'r').read()
                os.chdir("..")


                if hipdm_val == "1":
                     
                    top = [
                            ["Date&Time: " + formatted_date , "" ,""],
                            ["Name:" + patient_name , "Age:" + patient_age , "Gender:" + patient_gender],
                            ["Address:" + patient_address , "" , ""],
                            ["Date Of Admission:"+ add_date  , "Date of Discharge:" + dis_date , ""]


                        ]
                

                else:
                    top = [
                            ["Date&Time: " + formatted_date , "" ,""],
                            ["Name:" + patient_name , "Age:" + patient_age , "Gender:" + patient_gender],
                            ["Address:" + patient_address , "" , ""],
                           # ["Date Of Admission:"+ add_date +  "Date of Discharge:" + dis_date , ""]


                        ]
                







                    

                table1 = doc.add_table(rows=1 , cols=3)


                    
                for one , two , three in top:
                        cells1 = table1.add_row().cells
                        cells1[0].text = one
                        cells1[1].text = two
                        cells1[2].text = three

                table1.style = 'Light List'






                os.chdir("settings")
                inv_title = open("invoice_title.set" , 'r').read()
                os.chdir("..")
                doc.add_heading(str(inv_title), 2)




                table_header = ["SL.NO" , "Item" , "Quantity" , "Price"]

                table = doc.add_table(rows=5 , cols=5)

                for i in range(4):
                        table.rows[0].cells[i].text = table_header[i]

                sl_count = 0
                global tp
                for service , quantiry ,pri , gst in res:
                        int(sl_count)
                        sl_count +=1
                        cells = table.add_row().cells
                        cells[0].text = str(sl_count)
                        cells[1].text = service
                        cells[2].text = pri
                        cells[3].text=  str(  "Rs." +   quantiry)

                cells6 = table.add_row().cells
                cells6[0].text='-'
                cells6[1].text='-'
                cells6[2].text='-'
                cells6[3].text= "Total:- Rs." +str(tp)





                table.style = 'Colorful List'


                doc.add_heading("Grand Total:- Rs." + str(tp) , 4)

                doc.add_heading("Amount in Words: Rupees " + number_to_word(tp) + " Only." , 4)






                


                table.style = 'Colorful List'



                #doc.add_heading("Amount in Words: " + number_to_word(total) + " Only." , 4)

                doc.add_paragraph("")
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                para = doc.add_paragraph('Authorized Signatory:_______________________')
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT


                os.chdir("bills_created")

                os.remove(patient_name + "_bill.docx")
                
                # Now save the document to a location
                doc.save(patient_name + "_bill.docx")

                from tkinter import messagebox



                cur4.execute("DELETE FROM bill_data WHERE name=?", (name,))
                
                
                cur4.execute("DELETE FROM bill_data_cust WHERE name=?", (name,))




                cur4.execute("INSERT INTO bill_data VALUES (? , ? , ? , ? , ? , ? )",(patient_name, str(res),tp , patient_gender , patient_address , patient_age))
                
                print(formatted_date)

                cur4.execute("INSERT INTO bill_data_cust VALUES (? , ? , ? , ?)",(patient_name,patient_address,tp,formatted_date))

                con4.commit()
                con4.close()
                messagebox.showinfo("" , 'SUCESSFULLY EDITED AND SAVE INVOICE' , parent=win_edit)
                import webbrowser

                webbrowser.open(patient_name +"_bill.docx" )
                
                os.chdir("..")
                
                win_edit.destroy()
                import main
                refresh()






























            

            def addItem():
                global treev_win1
                #getting
                item_name = manual_itemname_entry_entry_win1.get()

                if item_name == "":
                    item_name = clicked1.get()

                item_price = manual_itemprice_entry_entry_win1.get()


                item_quantity = itemquantity_entry_entry_win1.get()

                item_gst = "12%"

                global tp
                tp += int(item_price)




                

                res.append([item_name , item_price , item_quantity , item_gst])
                
                
                treev2.insert("", 'end', text ="L1", 
                        values = (item_name , item_price , item_quantity, item_gst))
                


                manual_itemname_entry_entry_win1.delete(0, tk.END)
                

                manual_itemprice_entry_entry_win1.delete(0,tk.END)

                itemquantity_entry_entry_win1.delete(0, tk.END)

                show_total1.config(text="TOTAL=" + str(tp))

                print(res)
























            def remove_from_item_list_view_win_edit():
                global treev2

                slecteditem_win1_item_list_view = treev2.focus()


                details = treev2.item(slecteditem_win1_item_list_view)

                global total_price_count_win_edit
                print(type(total_price_count_win_edit))
                print(total_price_count_win_edit)
                print(type(total_price_count_win_edit))

                global tp
                tp -= int(details.get("values")[1])




                res.remove([str(details.get("values")[0])   , str(details.get("values")[1]) , str( details.get("values")[2])    ,   str(details.get("values")[3] ) ])


                treev2.delete(slecteditem_win1_item_list_view)
                show_total1.config(text="TOTAL=" + str(tp))







                       
            manual_itemname_entry_text_win1 = tk.Label(win_edit , text="ENTER ITEM NAME MANUALLY:-")
            manual_itemname_entry_text_win1.pack(anchor="n" , side="left" , padx=2 , pady=10)
            
            manual_itemname_entry_entry_win1 = tk.Entry(win_edit)
            manual_itemname_entry_entry_win1.pack(anchor="n" , side="left" , pady=10)
            
            
            
            
            
            
            
            
            manual_itemprice_entry_text_win1 = tk.Label(win_edit , text="ENTER ITEM PRICE MANUALLY:-")
            manual_itemprice_entry_text_win1.pack(anchor="n" , side="left" , padx=10 , pady=10)
            
            manual_itemprice_entry_entry_win1 = tk.Entry(win_edit)
            manual_itemprice_entry_entry_win1.pack(anchor="n" , side="left" , pady=10)
            
            
            
            
            
            itemquantity_entry_text_win1 = tk.Label(win_edit , text="ENTER Quantity:-")
            itemquantity_entry_text_win1.pack(anchor="n" , side="left" , padx=10 , pady=10)
            
            itemquantity_entry_entry_win1 = tk.Entry(win_edit)
            itemquantity_entry_entry_win1.pack(anchor="n" , side="left" , pady=10)


            #DROP DOWN


            cur2.execute("SELECT * FROM item")

            get_data_for_drop_down_menu = cur2.fetchall()

            options = []
            
            for sorted_get_data_drop in get_data_for_drop_down_menu:
                options.append(sorted_get_data_drop[0])


    
    # datatype of menu text 
            clicked1 = tk.StringVar() 
            
            # initial menu text 

            
            # Create Dropdown menu 
            drop1_win1 = tk.OptionMenu( win_edit , clicked1 , *options ) 

            or_other_optin_label_win1 = tk.Label(win_edit , text="OR select from list of added items:-")
            or_other_optin_label_win1.place(x=1 , y=47)

            drop1_win1.place(x=200 , y=47)


                        
            add_item_button_win1 = tk.Button(win_edit , text="ADD ITEM ✅" , bootstyle="outlined" , command=addItem)
            add_item_button_win1.place(x=20 , y=390)


            
            os.chdir("settings")
            hpval = open("hospital_ipd_mode.set" , 'r').read()
            os.chdir("..")


            if hpval == "1":
                 
                cal_addmission_text = tk.Label(win_edit , text="Date Of Addmission:-" )
                cal_addmission_text.place(x=10 , y=650)
                cal_addmission = tk.DateEntry(win_edit , bootstyle="danger")
                cal_addmission.place(x=100 , y=650)



                cal_dis_text = tk.Label(win_edit , text="Date Of Discharge:-" )
                cal_dis_text.place(x=470 , y=650)
                cal_dis = tk.DateEntry(win_edit , bootstyle="danger")
                cal_dis.place(x=560 , y=650)


            save_bill_win1 = tk.Button(win_edit , text="SAVE BILL/INVOICE 💾" , bootstyle="outlined" , command=save_bill_win_edit_function)
            save_bill_win1.place(x=970 , y=390)

            
            global show_total1

            show_total1 = tk.Label(win_edit , text="TOTAL="+str(total_price_count_win_edit) , font=("Ariel" , 16) , bootstyle = "danger")
            show_total1.place(x=980 , y=300)




                    















































            for ixt in res:
                treev2.insert("", 'end', text ="L1", 
                        values = (ixt[0] , ixt[1] , ixt[2] , ixt[3]))
                
            



                                








            remove_from_item_list_button_win_edit = tk.Button(win_edit , text="REMOVE AN SELECTED ITEM" , command=remove_from_item_list_view_win_edit)
            remove_from_item_list_button_win_edit.place(x=10 , y=300)
                    







            win_edit.mainloop()





















            

            









def open_bill():
             
             global treev
             
             slecteditem_win1_item_list_view = treev.focus()


             details = treev.item(slecteditem_win1_item_list_view)

             import webbrowser

             os.chdir("bills_created")

             webbrowser.open(str(details.get("values")[0] + "_bill.docx"))
     
     














edit_bill_launcher_button = tk.Button(tab2, text="Edit Selected Bill 🧑‍💻" ,bootstyle="success.Outline.TButton", command=edit_bill)
edit_bill_launcher_button.pack()

oepn_bill_launcher_button = tk.Button(tab2, text="Open Selected Bill 🧑‍💻" ,bootstyle="success.Outline.TButton", command=open_bill)
oepn_bill_launcher_button.pack()

#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------#











#item Management:

add_an_item_to_database_button = tk.Button(tab3 , text="Add an item/service to database" , command=this_app_pkg.item_manager.add_an_item_to_database_window)
add_an_item_to_database_button.pack(pady=10)

view_all_item_from_data_base_button = tk.Button(tab3 , text="View all items in database" , command=this_app_pkg.item_manager.view_item_window)
view_all_item_from_data_base_button.pack(pady=10)


#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------#





#SETTINGS:


from tkinter import Tk     # from tkinter import Tk for Python 3.x
from tkinter.filedialog import askopenfilename


def change_com_name():
    from tkinter import messagebox
    from tkinter.simpledialog import askstring
    

    
    
    com_name = askstring('////\\\\', 'Enter Your Company name:')
    os.chdir("settings")
    with open("com_name.set" , 'w')as f:
        f.write(com_name)
        f.close()
    os.chdir("..")

    messagebox.showinfo("SUCESS!" , "COMPANY NAME UPDATED SUCESSFULLY!                                                   RESTART APP TO MAKE CHANGES TO TAKE PLACE")





def change_logo_settings():

    filename = askopenfilename() # show an "Open" dialog box and return the path to the selected file
    os.chdir("settings")

    with open("logo_path.set" , 'w')as f:
        f.write(str(filename))
        f.close()



    os.chdir("..")

    from tkinter import messagebox
    messagebox.showinfo('SUCESS!',"LOGO SAVED SUCESSFULLY")


def change_invoice_title():
    from tkinter import messagebox
    from tkinter.simpledialog import askstring
    

    
    
    title_name_cache = askstring('////\\\\', 'Enter Invoice TITLE:-')

    os.chdir("settings")

    with open("invoice_title.set" , 'w')as f:
        f.write(str(title_name_cache))
        f.close()

    os.chdir("..")
    messagebox.showinfo('SUCESS!',"SETTING SAVE SUCESSFULLY!!!")


def set_hospital_ipd_mode():
    print("RAN")

    if messagebox.askyesnocancel("" , "Do you want to HOSPITAL IPD MODE to BE ON?           CLICK [YES] to turn {ON} and CLICK [NO] to turn {OFF}") == True:

        
        os.chdir("settings")
        with open("hospital_ipd_mode.set" , 'w')as f:
            f.write("1")
            f.close()

        os.chdir("..")
        messagebox.showinfo("" , "SET SUCESSFULLY")

    else:
        os.chdir("settings")
        with open("hospital_ipd_mode.set" , 'w')as f:
            f.write("0")
            f.close()

        os.chdir("..")
        messagebox.showinfo("SUCESS!" , "SETTING SAVED SUCESSFULLY!!")






change_company_name_button_tab5 = tk.Button(tab5 , text="CHANGE COMPANY NAME" , bootstyle="success.Outline.TButton" , command=change_com_name)
change_company_name_button_tab5.place(x=10 , y=20)


change_invoice_title_button_tab5 = tk.Button(tab5 , text="CHANGE INVOICE TITLE" , bootstyle="success.Outline.TButton" , command=change_invoice_title)
change_invoice_title_button_tab5.place(x=10 , y=50)



change_invoice_logo_button_tab5 = tk.Button(tab5 , text="CHANGE INVOICE LOGO" , bootstyle="success.Outline.TButton" , command=change_logo_settings)
change_invoice_logo_button_tab5.place(x=10 , y=80)


change_hospital_ipd_mode_button_tab5 = tk.Button(tab5 , text="SET HOSPITAL IPD MODE ON OR OFF" , bootstyle="success.Outline.TButton" , command=set_hospital_ipd_mode)
change_hospital_ipd_mode_button_tab5.place(x=10 , y=110)


















#----------------------------------------------------------------------------------------------------------------------------------------------------------------------------#































#LOOP rootS:
get_time()
root.mainloop()






