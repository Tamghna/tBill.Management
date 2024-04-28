import ttkbootstrap as tk
import sqlite3
import sqlite3
import ttkbootstrap as tk 
import tkinter as otk 
import datetime
import time
import os
import tsoft_basic_pkg
import this_app_pkg
from tkinter import ttk 
from decimal import *









try:

    class bill_manager():
        def create_bill_window():

            def number_to_word(number):
                def get_word(n):
                    words={ 0:"", 1:"One", 2:"Two", 3:"Three", 4:"Four", 5:"Five", 6:"Six", 7:"Seven", 8:"Eight", 9:"Nine", 10:"Ten", 11:"Eleven", 12:"Twelve", 13:"Thirteen", 14:"Fourteen", 15:"Fifteen", 16:"Sixteen", 17:"Seventeen", 18:"Eighteen", 19:"Nineteen", 20:"Twenty", 30:"Thirty", 40:"Forty", 50:"Fifty", 60:"Sixty", 70:"Seventy", 80:"Eighty", 90:"Ninty" }
                    if n<=20:
                        return words[n]
                    else:
                        ones=n%10
                        tens=n-ones
                        return words[tens]+" "+words[ones]
                        
                def get_all_word(n):
                    d=[100,10,100,100]
                    v=["","Hundred And","Thousand","lakh"]
                    w=[]
                    for i,x in zip(d,v):
                        t=get_word(n%i)
                        if t!="":
                            t+=" "+x
                        w.append(t.rstrip(" "))
                        n=n//i
                    w.reverse()
                    w=' '.join(w).strip()
                    if w.endswith("And"):
                        w=w[:-3]
                    return w

                arr=str(number).split(".")
                number=int(arr[0])
                crore=number//10000000
                number=number%10000000
                word=""
                if crore>0:
                    word+=get_all_word(crore)
                    word+=" crore "
                word+=get_all_word(number).strip()+" Rupees"
                if len(arr)>1:
                    if len(arr[1])==1:
                        arr[1]+="0"
                    word+=" and "+get_all_word(int(arr[1]))+" paisa"
                return word




            def save_bill_win1_function():

                con2 = sqlite3.connect("bill_data.db")
                cur2 = con2.cursor()

                #getting

                import time
                from time import strftime

                formatted_date = strftime("%d-%m-%Y %H:%M:%S")
        

                patient_name = cust_name_entry_win1.get()
                patient_age = cust_age_entry_win1.get()
                patient_address = cust_address_entry_win1.get()
                patient_gender = clicked2.get()




                os.chdir("settings")
                hpval = open("hospital_ipd_mode.set" , 'r').read()
                os.chdir("..")


                if hpval == "1":
                    add_date = cal_addmission.entry.get()
                    dis_date = cal_dis.entry.get()


  
                from docx import Document


                    # Create an instance of a word document
                doc = Document()



                    
                from docx.shared import Inches, Cm

                section = doc.sections[0]
                section.left_margin = Cm(1.0)
                section.right_margin = Cm(1.0)
                section.top_margin = Cm(1.0)
                

                try:
                     
                    os.chdir("settings")
                    logo_path = open("logo_path.set" , 'r').read()
                    doc.add_picture(str(logo_path), width=Inches(2), height=Inches(2))
                    os.chdir("..")
                except Exception:
                     from tkinter import messagebox
                     os.chdir("..")
                     messagebox.showerror("ERROR:00x120" , "AN ERROR OCCOURED.  ERROR:THE LOGO FILE IS NOT FOUND , IT MIGHT BE DELETED OR MOVED . PLEASE ASSIGN A NEW IMAGE FILE FROM THIS APPS' Settings TAB")
                     win1.destroy()
                     exit()




                os.chdir("settings")
                hipdm_val = open("hospital_ipd_mode.set" , 'r').read()
                os.chdir("..")


                if hipdm_val == "1":
                     
                    top = [
                            ["Date&Time: " + formatted_date , "" ,""],
                            ["Name:" + patient_name , "Age:" + patient_age , "Gender:" + patient_gender],
                            ["Address:" + patient_address , "" , ""],
                            ["Date Of Admission:"+ add_date  , "Date of Discharge:" + dis_date , ""]


                        ]
                

                else:
                    top = [
                            ["Date&Time: " + formatted_date , "" ,""],
                            ["Name:" + patient_name , "Age:" + patient_age , "Gender:" + patient_gender],
                            ["Address:" + patient_address , "" , ""],
                           # ["Date Of Admission:"+ add_date +  "Date of Discharge:" + dis_date , ""]


                        ]
                







                    

                table1 = doc.add_table(rows=1 , cols=3)


                    
                for one , two , three in top:
                        cells1 = table1.add_row().cells
                        cells1[0].text = one
                        cells1[1].text = two
                        cells1[2].text = three

                table1.style = 'Light List'






                os.chdir("settings")
                inv_title = open("invoice_title.set" , 'r').read()
                os.chdir("..")
                doc.add_heading(str(inv_title), 2)




                table_header = ["SL.NO" , "Item" , "Quantity" , "Price"]

                table = doc.add_table(rows=4 , cols=4)

                for i in range(4):
                        table.rows[0].cells[i].text = table_header[i]

                sl_count = 0

                for service , quantiry ,pri , gst in cached_list_items_win1:
                        int(sl_count)
                        sl_count +=1
                        cells = table.add_row().cells
                        cells[0].text = str(sl_count)
                        cells[1].text = service
                        cells[2].text = pri
                        cells[3].text=  str(  "Rs." +   quantiry)
                
                cells5 = table.add_row().cells
                cells5[0].text=''
                cells5[1].text=''
                cells5[2].text=''
                cells5[3].text= "Total:- Rs." +str(total_price_count)





                table.style = 'Colorful List'


                doc.add_heading("Grand Total:- Rs." + str(total_price_count) , 4)

                doc.add_heading("Amount in Words: " + number_to_word(total_price_count) + " Only." , 4)




                doc.add_paragraph("")
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                para = doc.add_paragraph('Authorized Signatory:_______________________')
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT


                os.chdir("bills_created")
                
                # Now save the document to a location
                doc.save(patient_name + "_bill.docx")

                from tkinter import messagebox


                cur2.execute("INSERT INTO bill_data VALUES (? , ? , ? , ? , ? , ? )",(patient_name, str(cached_list_items_win1),total_price_count , patient_gender , patient_address , patient_age))
                
                print(formatted_date)

                cur2.execute("INSERT INTO bill_data_cust VALUES (? , ? , ? , ?)",(patient_name,patient_address,total_price_count,formatted_date))

                con2.commit()
                con2.close()
                messagebox.showinfo("" , 'SUCESSFULLY SAVED BILL/INVOICE' , parent=win1)
                import webbrowser
                webbrowser.open(patient_name +"_bill.docx" )
                
                os.chdir("..")
                
                win1.destroy()
                import main
                main.refresh()

                
                

    
























            cached_list_items_win1 = []
            global total_price_count
            total_price_count = 0

            




            def addItem():

                global treev_win1
                #getting
                item_name = manual_itemname_entry_entry_win1.get()

                if item_name == "":
                    item_name = clicked1.get()

                item_price = manual_itemprice_entry_entry_win1.get()


                item_quantity = itemquantity_entry_entry_win1.get()

                item_gst = "12%"

                global total_price_count


                total_price_count += int(item_price)




                

                cached_list_items_win1.append([item_name , item_price , item_quantity , item_gst])
                
                
                treev_win1.insert("", 'end', text ="L1", 
                        values = (item_name , item_price , item_gst , item_quantity))
                
                print(cached_list_items_win1)
                print(total_price_count)


                manual_itemname_entry_entry_win1.delete(0, tk.END)
                

                manual_itemprice_entry_entry_win1.delete(0,tk.END)

                itemquantity_entry_entry_win1.delete(0, tk.END)

                show_total.config(text="TOTAL=" + str(total_price_count))







            def remove_from_item_list_view_win1():
                global treev_win1

                slecteditem_win1_item_list_view = treev_win1.focus()


                details = treev_win1.item(slecteditem_win1_item_list_view)


                global total_price_count



                r_itm = details.get("values")[1]



                total_price_count -= int(r_itm)


                cached_list_items_win1.remove([str(details.get("values")[0])   , str(details.get("values")[1]) , str( details.get("values")[3])    ,   str(details.get("values")[2] ) ])


                treev_win1.delete(slecteditem_win1_item_list_view)


                


                
                show_total.config(text="TOTAL=" + str(total_price_count))
                print(cached_list_items_win1)
                print(total_price_count)















            con1 = sqlite3.connect("items.db")
            cur1 = con1.cursor()


            win1 = tk.Toplevel()
            win1.attributes('-fullscreen', True)
            win1.geometry("1300x800")
            win1.title("CREATE A NEW BILL _ WINDOW")
            win1.config(background="#b3c6ff")
            from tkinter import messagebox

            def exit_command():
                 if messagebox.askyesnocancel("" , "WARNING :THIS ACTION WILL EXIT THIS WINDOW AND GO BACK TO THE HOME SCREEN .. ANY UNSAVED WORK WILL BE LOST!  .. DO YOU WANT TO EXIT?" , parent=win1):
                      win1.destroy()

            exit_button = tk.Button(win1 , text="EXIT❌" , command=exit_command)
            exit_button.place(x=1200 , y=10)            
            
            heading_win1 = tk.Label(win1,text="Create New Bill/Invoice :-" , font=("London-Tube" , 19))
            heading_win1.pack(pady=10)
            
            
            manual_itemname_entry_text_win1 = tk.Label(win1 , text="ENTER ITEM NAME MANUALLY:-")
            manual_itemname_entry_text_win1.pack(anchor="n" , side="left" , padx=2 , pady=10)
            
            manual_itemname_entry_entry_win1 = tk.Entry(win1)
            manual_itemname_entry_entry_win1.pack(anchor="n" , side="left" , pady=10)
            
            
            
            
            
            
            
            
            manual_itemprice_entry_text_win1 = tk.Label(win1 , text="ENTER ITEM PRICE MANUALLY:-")
            manual_itemprice_entry_text_win1.pack(anchor="n" , side="left" , padx=10 , pady=10)
            
            manual_itemprice_entry_entry_win1 = tk.Entry(win1)
            manual_itemprice_entry_entry_win1.pack(anchor="n" , side="left" , pady=10)
            
            
            
            
            
            itemquantity_entry_text_win1 = tk.Label(win1 , text="ENTER Quantity:-")
            itemquantity_entry_text_win1.pack(anchor="n" , side="left" , padx=10 , pady=10)
            
            itemquantity_entry_entry_win1 = tk.Entry(win1)
            itemquantity_entry_entry_win1.pack(anchor="n" , side="left" , pady=10)


            #DROP DOWN


            cur1.execute("SELECT * FROM item")

            get_data_for_drop_down_menu = cur1.fetchall()

            options = []
            
            for sorted_get_data_drop in get_data_for_drop_down_menu:
                options.append(sorted_get_data_drop[0])


    
    # datatype of menu text 
            clicked1 = tk.StringVar() 
            
            # initial menu text 

            
            # Create Dropdown menu 
            drop1_win1 = tk.OptionMenu( win1 , clicked1 , *options ) 

            or_other_optin_label_win1 = tk.Label(win1 , text="OR select from list of added items:-")
            or_other_optin_label_win1.place(x=1 , y=93)

            drop1_win1.place(x=200 , y=93)
                    
                
            
            global treev_win1
            
            treev_win1 = ttk.Treeview(win1, selectmode ='browse')

            
    
            # Calling pack method w.r.to treeview
            treev_win1.place(x=20 , y=150)
            
            # Constructing vertical scrollbar
            # with treeview
            verscrlbar = ttk.Scrollbar(win1, 
                                    orient ="vertical", 
                                    command = treev_win1.yview)
            
            # Calling pack method w.r.to vertical 
            # scrollbar
            verscrlbar.pack()
            
            # Configuring treeview
            treev_win1.configure(xscrollcommand = verscrlbar.set)
            
            # Defining number of columns
            treev_win1["columns"] = ("1", "2", "3" , "4")
            
            # Defining heading
            treev_win1['show'] = 'headings'
            
            # Assigning the width and anchor to  the
            # respective columns
            treev_win1.column("1", width = 270, anchor ='c')
            treev_win1.column("2", width = 150, anchor ='c')
            treev_win1.column("3", width = 150, anchor ='c')
            treev_win1.column("4", width = 150, anchor ='c')
            
            # Assigning the heading names to the 
            # respective columns
            treev_win1.heading("1", text ="ITEM NAME")
            treev_win1.heading("2", text ="PRICE")
            treev_win1.heading("3", text ="GST")
            treev_win1.heading("4", text ="QUANTITY")
            
            # Inserting the items and their features to the 
            # columns built




                    
            
            
            add_item_button_win1 = tk.Button(win1 , text="ADD ITEM ✅" , bootstyle="outlined" , command=addItem)
            add_item_button_win1.place(x=20 , y=390)
            
            remove_selected_item_button_win1 = tk.Button(win1 , text="REMOVE AN SELECTED ITEM ❌" , bootstyle="outlined" , command=remove_from_item_list_view_win1)
            remove_selected_item_button_win1.place(x=170 , y=390)

            save_bill_win1 = tk.Button(win1 , text="SAVE BILL/INVOICE 💾" , bootstyle="outlined" , command=save_bill_win1_function)
            save_bill_win1.place(x=970 , y=390)
            
            save_bill_and_print_win1 = tk.Button(win1 , text="SAVE BILL AND PRINT 💾🖨️" , bootstyle="outlined")
            save_bill_and_print_win1.place(x=970 , y=420)
            
            
            
            fill_this_details_text_win1 = tk.Label(win1 , text="Fill the Patient Data before adding items!:-" , font=("London-Tube" , 15 ) , background="red")
            fill_this_details_text_win1.place(x=10 , y=450)
            
            cust_name_text_win1 = tk.Label(win1 , text="Patient Name:")
            cust_name_text_win1.place(x=10 , y=490)
            cust_name_entry_win1 = tk.Entry(win1)
            cust_name_entry_win1.place(x=119 , y=490)
            
            
            cust_age_text_win1 = tk.Label(win1 , text="Patient Age:")
            cust_age_text_win1.place(x=10 , y=515)
            cust_age_entry_win1 = tk.Entry(win1)
            cust_age_entry_win1.place(x=119 , y=515)
            
            
            
            cust_address_text_win1 = tk.Label(win1 , text="Patient Address:")
            cust_address_text_win1.place(x=10 , y=542)
            cust_address_entry_win1 = tk.Entry(win1)
            cust_address_entry_win1.place(x=120  , y=542)


            global show_total

            show_total = tk.Label(win1 , text="Total=0" , font=("Ariel" , 16) , bootstyle = "danger")
            show_total.place(x=980 , y=300)
            
            
            


            clicked2 = tk.StringVar()
            options2=["----SELECT GENDER----" , "MALE" , "FEMALE"]
            cust_gender_text_win1 = tk.Label(win1 , text="Patient Gender:")
            cust_gender_text_win1.place(x=10 , y=574)
            clicked2 = tk.StringVar() 
            
            # initial menu text 
            clicked2.set( "----SELECT GENDER----" ) 
            
            # Create Dropdown menu 
            drop2_win1 = tk.OptionMenu( win1 , clicked2 , *options2 ) 


            drop2_win1.place(x=137 , y=574)



            os.chdir("settings")
            hpval = open("hospital_ipd_mode.set" , 'r').read()
            os.chdir("..")


            if hpval == "1":
                 
                cal_addmission_text = tk.Label(win1 , text="Date Of Addmission:-" )
                cal_addmission_text.place(x=10 , y=650)
                cal_addmission = tk.DateEntry(win1 , bootstyle="danger")
                cal_addmission.place(x=100 , y=650)



                cal_dis_text = tk.Label(win1 , text="Date Of Discharge:-" )
                cal_dis_text.place(x=470 , y=650)
                cal_dis = tk.DateEntry(win1 , bootstyle="danger")
                cal_dis.place(x=560 , y=650)

            con1.close()



            
            
            
            
            
            
            
            
            win1.mainloop()
            
        
        
        
        
        
        def edit_bill_window():
            win4 = tk.Toplevel()
            win4.geometry("1300x800")
            win4.title("EDIT BILL   |   tBillManager")

            
            
            
            
            
            
            
            
            win4.geometry()








    #************************************************************************************************************************************************************#



    class item_manager():
        def add_an_item_to_database_window():
            
            
            def add_exe():
                from tkinter import messagebox
                import os
                print(os.curdir)
                item_name_get = win2_item_name_entry.get()
                item_price_get = win2_item_price_entry.get()
                item_gst_get = win2_item_gst_entry.get()
                
                conn2 = sqlite3.connect("items.db")
                cur2 = conn2.cursor()

                cur2.execute("insert into item values (?, ?, ?)",(item_name_get, item_price_get, item_gst_get))
                
                conn2.commit()
                
                win2.destroy()
                
                messagebox.showinfo("" , "SUCESS!")
                
            
            
            
            
            
            
            
            
            
            win2 = tk.Toplevel()
            win2.geometry("1200x600")
            win2.title("Add item to database")
            
            win2_item_name_text = tk.Label(win2 , text="Item name:-")
            win2_item_name_text.pack()
            
            win2_item_name_entry = tk.Entry(win2)
            win2_item_name_entry.pack()
            
            
            
            win2_item_price_text = tk.Label(win2 , text="Item Price:-")
            win2_item_price_text.pack()
            
            win2_item_price_entry = tk.Entry(win2)
            win2_item_price_entry.pack(padx=5)
            
            
            
            win2_item_gst_text = tk.Label(win2 , text="Item GST%:-")
            win2_item_gst_text.pack()
            
            win2_item_gst_entry = tk.Entry(win2)
            win2_item_gst_entry.pack(padx=5)
            

            
            
            win2_save_button = tk.Button(win2 , text="SAVE" , command=add_exe)
            win2_save_button.pack(anchor="center")
            
            
            
            win2.mainloop()
            
            
            
            
        def view_item_window():
            global lcount
            lcount = 0

            from tkinter import ttk 

            win3 = tk.Window()
            win3.geometry("1000x400")
            
            conn2 = sqlite3.connect("items.db")
            cur2 = conn2.cursor()
            
            
            
            treev = ttk.Treeview(win3, selectmode ='browse')
    
            # Calling pack method w.r.to treeview
            treev.pack()
            
            # Constructing vertical scrollbar
            # with treeview
            verscrlbar = ttk.Scrollbar(win3, 
                                    orient ="vertical", 
                                    command = treev.yview)
            
            # Calling pack method w.r.to vertical 
            # scrollbar
            verscrlbar.pack(side="right",fill ='x')
            
            # Configuring treeview
            treev.configure(xscrollcommand = verscrlbar.set)
            
            # Defining number of columns
            treev["columns"] = ("1", "2", "3" )
            
            # Defining heading
            treev['show'] = 'headings'
            
            # Assigning the width and anchor to  the
            # respective columns
            treev.column("1", width = 230, anchor ='c')
            treev.column("2", width = 150, anchor ='se')
            treev.column("3", width = 150, anchor ='se')

            
            # Assigning the heading names to the 
            # respective columns
            treev.heading("1", text ="Item Name")
            treev.heading("2", text ="Price in ₹")
            treev.heading("3", text ="GST%")

            # Inserting the items and their features to the 
            # columns built
            cur2.execute("SELECT * FROM item")

            data = cur2.fetchall()

            for single in data:
                one = single[0]
                two = single[1]
                three= single[2]
                
                lcount +=1
                
                
                treev.insert("", 'end', text = "L" + str(lcount), 
                            values = (one,two,three))
                    
                

                
                


                    
                    
                    
                    
                    
            
            
            
            
            
            
            
            
            win3.mainloop()
            

except Exception as e:
    from tkinter import messagebox
    messagebox.showerror("" , "AN ERROR OCCOURED!       ERROR IS :- " + str(e))()
